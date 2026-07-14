from __future__ import annotations

import json
import os
from collections.abc import Iterator
from contextlib import contextmanager
from pathlib import Path, PurePosixPath
from typing import Literal

from docx import Document
from openpyxl import Workbook  # type: ignore[import-untyped]
from pydantic import Field

from . import __version__
from .config import create_repository, repository_config_path
from .engine import UpdateEngine
from .models import OctopusModel, utc_now
from .rendering import read_machine_header
from .search import SEARCH_ALGORITHM_VERSION, SearchIndex, analyze_terms
from .utils import atomic_write_json, atomic_write_text, load_json

REQUIRED_SEARCH_CATEGORIES = {"chinese", "english", "text", "non_text", "same_name", "stale"}
TOP5_THRESHOLD = 0.80
MRR_THRESHOLD = 0.65
INSPECTION_REDUCTION_THRESHOLD = 0.30


def default_search_evaluation_dataset_path() -> Path:
    packaged = Path(__file__).resolve().parent / "data" / "search-value-v1.json"
    if packaged.exists():
        return packaged
    return Path(__file__).resolve().parents[2] / "benchmarks" / "datasets" / "search-value-v1.json"


class EvaluationDocument(OctopusModel):
    path: str
    format: Literal["docx", "xlsx"]
    title: str
    content: list[str] = Field(default_factory=list)
    stale: bool = False


class EvaluationTask(OctopusModel):
    task_id: str
    query: str
    relevant_paths: list[str]
    categories: list[str]


class SearchEvaluationDataset(OctopusModel):
    schema_version: Literal["1.0"]
    dataset_id: str
    dataset_version: str
    required_categories: list[str]
    documents: list[EvaluationDocument]
    tasks: list[EvaluationTask]


class SearchEvaluationTaskResult(OctopusModel):
    task_id: str
    query: str
    categories: list[str]
    relevant_paths: list[str]
    returned_paths: list[str]
    search_rank: int | None = None
    filename_baseline_rank: int
    top5_hit: bool
    reciprocal_rank: float
    inspection_steps_saved: int
    explanation_contract_pass: bool
    failure_reason: str = ""


class SearchEvaluationReport(OctopusModel):
    product_version: str
    search_algorithm_version: str
    dataset_id: str
    dataset_version: str
    generated_at: str
    task_count: int
    top5_accuracy: float
    mean_reciprocal_rank: float
    task_failure_count: int
    mean_inspection_step_reduction: float
    explanation_contract_failures: int
    thresholds: dict[str, float]
    meets_engineering_thresholds: bool
    tasks: list[SearchEvaluationTaskResult]


def load_search_evaluation_dataset(path: Path) -> SearchEvaluationDataset:
    dataset = SearchEvaluationDataset.model_validate(load_json(path))
    validate_search_evaluation_dataset(dataset)
    return dataset


def validate_search_evaluation_dataset(dataset: SearchEvaluationDataset) -> None:
    if not dataset.dataset_id.strip() or not dataset.dataset_version.strip():
        raise ValueError("dataset_id and dataset_version are required")
    if not dataset.documents or not dataset.tasks:
        raise ValueError("evaluation dataset requires documents and tasks")
    paths: set[str] = set()
    for document in dataset.documents:
        pure = PurePosixPath(document.path)
        if pure.is_absolute() or ".." in pure.parts or not pure.name:
            raise ValueError(f"unsafe evaluation document path: {document.path}")
        normalized = pure.as_posix()
        if normalized in paths:
            raise ValueError(f"duplicate evaluation document path: {normalized}")
        if pure.suffix.casefold() != f".{document.format}":
            raise ValueError(f"document format does not match path: {normalized}")
        paths.add(normalized)
    task_ids: set[str] = set()
    covered: set[str] = set()
    for task in dataset.tasks:
        if task.task_id in task_ids:
            raise ValueError(f"duplicate evaluation task_id: {task.task_id}")
        task_ids.add(task.task_id)
        if not task.query.strip() or not task.relevant_paths:
            raise ValueError(f"task requires a query and relevant_paths: {task.task_id}")
        missing = set(task.relevant_paths) - paths
        if missing:
            raise ValueError(f"task {task.task_id} references missing paths: {sorted(missing)}")
        covered.update(task.categories)
    required = set(dataset.required_categories)
    if not required >= REQUIRED_SEARCH_CATEGORIES:
        missing_categories = sorted(REQUIRED_SEARCH_CATEGORIES - required)
        raise ValueError(f"dataset omits required categories: {missing_categories}")
    if not required <= covered:
        uncovered_categories = sorted(required - covered)
        raise ValueError(f"tasks do not cover required categories: {uncovered_categories}")


def _write_document(root: Path, document: EvaluationDocument) -> None:
    destination = root.joinpath(*PurePosixPath(document.path).parts)
    destination.parent.mkdir(parents=True, exist_ok=True)
    if document.format == "docx":
        value = Document()
        value.add_heading(document.title, level=1)
        for paragraph in document.content:
            value.add_paragraph(paragraph)
        value.save(str(destination))
        return
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = document.title[:31] or "Data"
    for row in document.content:
        sheet.append([cell.strip() for cell in row.split("|")])
    workbook.save(destination)


@contextmanager
def _isolated_config(root: Path) -> Iterator[None]:
    previous = {name: os.environ.get(name) for name in ("APPDATA", "XDG_CONFIG_HOME")}
    os.environ["APPDATA"] = str(root / "appdata")
    os.environ["XDG_CONFIG_HOME"] = str(root / "xdg-config")
    try:
        yield
    finally:
        for name, value in previous.items():
            if value is None:
                os.environ.pop(name, None)
            else:
                os.environ[name] = value


def _materialize_dataset(dataset: SearchEvaluationDataset, workspace: Path) -> Path:
    raw = workspace / "raw"
    index = workspace / "index"
    raw.mkdir(parents=True, exist_ok=True)
    for document in dataset.documents:
        _write_document(raw, document)
    config = create_repository(raw, index, f"Evaluation {dataset.dataset_id}", ai_enabled=False)
    config.stability.minimum_quiet_seconds = 0
    config.stability.required_stable_scan_count = 1
    atomic_write_json(
        repository_config_path(index),
        config.model_dump(mode="json", by_alias=True),
    )
    UpdateEngine(index).run(force_path="*")
    search = SearchIndex(index)
    by_source = {
        document.raw_relative_path: document
        for document in search._iter_documents()
        if document.index_type == "leaf"
    }
    missing = {
        document.path for document in dataset.documents if document.path not in by_source
    }
    if missing:
        raise RuntimeError(f"evaluation documents were not indexed: {sorted(missing)}")
    for definition in dataset.documents:
        if not definition.stale:
            continue
        indexed = by_source[definition.path]
        header, body = read_machine_header(Path(indexed.index_path))
        header.setdefault("update_control", {})["index_status"] = "stale"
        atomic_write_text(
            Path(indexed.index_path),
            json.dumps(header, ensure_ascii=False, indent=2) + "\n\n" + body,
        )
    search.rebuild()
    return index


def _filename_baseline_order(
    query: str, documents: list[EvaluationDocument]
) -> list[str]:
    query_terms = set(analyze_terms(query))
    scored: list[tuple[int, str]] = []
    for document in documents:
        name = PurePosixPath(document.path).name
        score = len(query_terms & set(analyze_terms(name)))
        if query.casefold().strip() in name.casefold():
            score += 100
        scored.append((score, document.path))
    scored.sort(key=lambda item: (-item[0], item[1].casefold()))
    return_paths = [path for _, path in scored]
    return return_paths


def evaluate_search_dataset(
    dataset: SearchEvaluationDataset,
    workspace: Path,
) -> SearchEvaluationReport:
    validate_search_evaluation_dataset(dataset)
    workspace = workspace.resolve()
    if workspace.exists() and any(workspace.iterdir()):
        raise ValueError(f"evaluation workspace must be empty: {workspace}")
    workspace.mkdir(parents=True, exist_ok=True)
    with _isolated_config(workspace):
        index = _materialize_dataset(dataset, workspace)
        search = SearchIndex(index)
        task_results: list[SearchEvaluationTaskResult] = []
        for task in dataset.tasks:
            results = search.search(task.query, limit=max(20, len(dataset.documents)))
            returned_paths = [item.raw_relative_path for item in results]
            relevant = set(task.relevant_paths)
            search_rank = next(
                (rank for rank, path in enumerate(returned_paths, start=1) if path in relevant),
                None,
            )
            filename_order = _filename_baseline_order(task.query, dataset.documents)
            filename_rank = next(
                rank
                for rank, path in enumerate(filename_order, start=1)
                if path in relevant
            )
            relevant_result = next(
                (item for item in results if item.raw_relative_path in relevant),
                None,
            )
            explanation_ok = bool(
                relevant_result
                and relevant_result.raw_relative_path
                and relevant_result.open_target_uri
                and relevant_result.match_evidence
                and relevant_result.evidence
            )
            if "stale" in task.categories:
                explanation_ok = explanation_ok and bool(
                    relevant_result and "stale_index" in relevant_result.risk_flags
                )
            top5_hit = search_rank is not None and search_rank <= 5
            effective_search_rank = search_rank or len(dataset.documents) + 1
            failure_reason = ""
            if not top5_hit:
                failure_reason = "relevant_document_not_in_top5"
            elif not explanation_ok:
                failure_reason = "explanation_contract_failed"
            task_results.append(
                SearchEvaluationTaskResult(
                    task_id=task.task_id,
                    query=task.query,
                    categories=task.categories,
                    relevant_paths=task.relevant_paths,
                    returned_paths=returned_paths[:5],
                    search_rank=search_rank,
                    filename_baseline_rank=filename_rank,
                    top5_hit=top5_hit,
                    reciprocal_rank=0.0 if search_rank is None else 1.0 / search_rank,
                    inspection_steps_saved=filename_rank - effective_search_rank,
                    explanation_contract_pass=explanation_ok,
                    failure_reason=failure_reason,
                )
            )
    task_count = len(task_results)
    top5_accuracy = sum(item.top5_hit for item in task_results) / task_count
    mrr = sum(item.reciprocal_rank for item in task_results) / task_count
    failures = sum(not item.top5_hit for item in task_results)
    contract_failures = sum(not item.explanation_contract_pass for item in task_results)
    baseline_steps = sum(item.filename_baseline_rank for item in task_results)
    search_steps = sum(
        item.search_rank if item.search_rank is not None else len(dataset.documents) + 1
        for item in task_results
    )
    reduction = 1.0 - (search_steps / max(1, baseline_steps))
    thresholds = {
        "top5_accuracy_min": TOP5_THRESHOLD,
        "mean_reciprocal_rank_min": MRR_THRESHOLD,
        "mean_inspection_step_reduction_min": INSPECTION_REDUCTION_THRESHOLD,
        "explanation_contract_failures_max": 0.0,
    }
    meets = (
        top5_accuracy >= TOP5_THRESHOLD
        and mrr >= MRR_THRESHOLD
        and reduction >= INSPECTION_REDUCTION_THRESHOLD
        and contract_failures == 0
    )
    return SearchEvaluationReport(
        product_version=__version__,
        search_algorithm_version=SEARCH_ALGORITHM_VERSION,
        dataset_id=dataset.dataset_id,
        dataset_version=dataset.dataset_version,
        generated_at=utc_now(),
        task_count=task_count,
        top5_accuracy=top5_accuracy,
        mean_reciprocal_rank=mrr,
        task_failure_count=failures,
        mean_inspection_step_reduction=reduction,
        explanation_contract_failures=contract_failures,
        thresholds=thresholds,
        meets_engineering_thresholds=meets,
        tasks=task_results,
    )
