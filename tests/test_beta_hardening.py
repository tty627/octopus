from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from octopus.engine import UpdateEngine
from octopus.parsers import ParserRegistry
from octopus.search import SearchIndex
from octopus.transactions import load_run_report
from octopus.utils import sha256_file


def test_unicode_path_beyond_legacy_windows_max_path_is_indexed_read_only(
    repository: tuple[Path, Path, object],
) -> None:
    raw, index, _ = repository
    directory = raw
    for number in range(7):
        directory /= f"第{number}层-同步资料-" + "长" * 18
    directory.mkdir(parents=True)
    source = directory / "最终证据.txt"
    source.write_text("LONGPATH-READY 可恢复长路径证据", encoding="utf-8")
    before = sha256_file(source)
    assert len(str(source.resolve())) > 260

    stats = UpdateEngine(index).run(force_path="*")
    result = next(
        item for item in SearchIndex(index).search("LONGPATH-READY") if item.name == source.name
    )

    assert stats.failed == 0
    assert load_run_report(index).status == "success"
    assert result.raw_relative_path.endswith("最终证据.txt")
    assert sha256_file(source) == before


def test_source_permission_failure_keeps_previous_index_and_reports_partial(
    repository: tuple[Path, Path, object], monkeypatch: pytest.MonkeyPatch
) -> None:
    raw, index, _ = repository
    source = raw / "permission.docx"
    document = Document()
    document.add_paragraph("previous searchable evidence")
    document.save(source)
    UpdateEngine(index).run(force_path="*")
    previous_result = next(
        item
        for item in SearchIndex(index).search("previous searchable")
        if item.name == source.name
    )
    previous_index = Path(previous_result.index_path)
    previous_bytes = previous_index.read_bytes()
    document.add_paragraph("new version that cannot be read yet")
    document.save(source)
    raw_before = source.read_bytes()

    def denied(self: ParserRegistry, path: Path) -> object:
        raise PermissionError(f"access denied for {path.name}")

    monkeypatch.setattr(ParserRegistry, "extract", denied)
    stats = UpdateEngine(index).run(force_path="permission.docx")
    report = load_run_report(index)

    assert stats.failed == 1
    assert report.status == "partial"
    assert report.errors[0]["code"] == "PermissionError"
    assert previous_index.read_bytes() == previous_bytes
    assert source.read_bytes() == raw_before
    assert SearchIndex(index).search("previous searchable")
